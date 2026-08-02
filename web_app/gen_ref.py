"""Generate reference DOCX v3 — issue description formatting"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from lxml import etree

desktop = os.path.expanduser('~/Desktop')
doc = Document()

style = doc.styles['Normal']; style.font.name = 'Microsoft YaHei'; style.font.size = Pt(10)
rPr = style.element.get_or_add_rPr()
rFonts = etree.SubElement(rPr, qn('w:rFonts')); rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

def set_cell(cell, text, bold=False, size=Pt(8.5), color=None):
    p = cell.paragraphs[0]; p.clear()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text).replace('**',''))
    run.font.name = 'Microsoft YaHei'; run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rf = etree.SubElement(rPr, qn('w:rFonts')); rf.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if bold: run.bold = True
    if color: run.font.color.rgb = color

def shade(cell, color='EEF2FF'):
    shd = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shd.set(qn('w:fill'), color); shd.set(qn('w:val'), 'clear')

cols = ['#','Type','Standard Requirement','Evidence from PPT','Cov','Score','Issue Description']
rows_a = [
    ['1.1','Key Result','Participate in company-level SP/BP formulation, guide product roadmap','Defined flashlight series user segmentation; clarified Baton series positioning & Baton 4 Pro combo strategy','⚠️','6/10','Evidence stays at product-line strategy level. Although quality is good, it addresses single product line only. S4 standard requires company brand-level SP/BP strategic participation — missing direct evidence of involvement in strategic planning meetings or roadmap documents.'],
    ['1.2','Key Behavior','Participate in company SP/BP strategic market review','In-depth analysis of all flashlight bestseller models; thinking holistically about persona & scenario differences','⚠️','6/10','Flashlight series user analysis shows some global perspective, but participation in SP/BP strategic review is not clearly demonstrated. Suggest adding specific meeting dates and contribution points to strategic review sessions.'],
    ['1.3','Key Behavior','Provide market recommendations for company product roadmap','Deep insight into Baton 4 Pro positioning via Five-Forces model; organized R&D deep-dive on selling points','⚠️','6/10','Product positioning and selling-point ranking work is solid, but roadmap recommendations require cross-product-line planning advice. Current evidence leans toward single-product tactical level.'],
]

# ============ VERSION A: 7-col flat table ============
doc.add_heading('Version A: Traditional 7-column flat table', level=2)
doc.add_paragraph('Every field in its own column. Issue description column gets 5cm — text wraps but still feels cramped because 7 columns is a lot.').runs[0].font.size = Pt(9)

t1 = doc.add_table(rows=len(rows_a)+1, cols=7); t1.style = 'Table Grid'; t1.alignment = 1; t1.autofit = False
w1 = [Cm(1.0), Cm(1.5), Cm(3.0), Cm(3.5), Cm(1.0), Cm(1.2), Cm(4.8)]
for ci, h in enumerate(cols):
    set_cell(t1.cell(0,ci), h, bold=True, size=Pt(8)); t1.cell(0,ci).width = w1[ci]; shade(t1.cell(0,ci))
for ri, row in enumerate(rows_a):
    for ci, val in enumerate(row):
        set_cell(t1.cell(ri+1,ci), val, size=Pt(8)); t1.cell(ri+1,ci).width = w1[ci]

doc.add_paragraph()

# ============ VERSION B: 6-col — merge std+evidence ============
doc.add_heading('Version B: 6-column — standard + evidence merged', level=2)
doc.add_paragraph('Standard requirement and employee evidence are shown together in one cell with line breaks. This frees up space for issue description to get a proper 6cm column.').runs[0].font.size = Pt(9)

cols_b = ['#','Type','Standard Requirement + Employee Evidence','Cov','Score','Issue Description']
rows_b = [
    ['1.1','Key Result','Standard: Participate in company-level SP/BP formulation, guide product roadmap.\nEvidence: Defined flashlight series user segmentation; clarified Baton series positioning & combo strategy; defined core selling points & priority ranking.','⚠️','6/10','Evidence stays at product-line strategy level. Although quality is good, it addresses single product line only. S4 standard requires company brand-level SP/BP strategic participation — missing direct evidence of involvement in strategic planning meetings or roadmap documents.'],
    ['1.2','Key Behavior','Standard: Participate in company SP/BP strategic market review.\nEvidence: In-depth analysis of all flashlight bestseller models across independent sites & Amazon; thinking holistically about persona & scenario differences.','⚠️','6/10','Flashlight series user analysis shows some global perspective, but participation in SP/BP strategic review is not clearly demonstrated. Suggest adding specific meeting dates and contribution points to strategic review sessions.'],
    ['1.3','Key Behavior','Standard: Provide market recommendations for company product roadmap.\nEvidence: Deep insight into Baton 4 Pro positioning via Five-Forces model; organized R&D deep-dive on selling points.','⚠️','6/10','Product positioning and selling-point ranking work is solid, but roadmap recommendations require cross-product-line planning advice. Current evidence leans toward single-product tactical level.'],
]

t2 = doc.add_table(rows=len(rows_b)+1, cols=6); t2.style = 'Table Grid'; t2.alignment = 1; t2.autofit = False
w2 = [Cm(1.0), Cm(1.5), Cm(5.5), Cm(1.0), Cm(1.2), Cm(5.8)]
for ci, h in enumerate(cols_b):
    set_cell(t2.cell(0,ci), h, bold=True, size=Pt(8)); t2.cell(0,ci).width = w2[ci]; shade(t2.cell(0,ci))
for ri, row in enumerate(rows_b):
    for ci, val in enumerate(row):
        set_cell(t2.cell(ri+1,ci), val, size=Pt(8))
        if ci == 5:  # issue description: use gray text
            p = t2.cell(ri+1,ci).paragraphs[0]
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            # add shading to issue cell
            shade(t2.cell(ri+1,ci), 'F9FAFB')
        t2.cell(ri+1,ci).width = w2[ci]

doc.add_paragraph()

# ============ VERSION C: Two-tier ============
doc.add_heading('Version C: Two-tier — compact summary row + detail row below', level=2)
doc.add_paragraph('Each assessment item gets a slim summary row, then a shaded full-width detail row underneath just for the issue description. Clean, scannable, annotation-friendly.').runs[0].font.size = Pt(9)

cols_c = ['#','Type','Standard Requirement','Evidence Summary','Cov','Score']
for idx, row_data in enumerate(rows_a):
    t = doc.add_table(rows=2, cols=6); t.style = 'Table Grid'; t.alignment = 1; t.autofit = False
    wc = [Cm(1.0), Cm(1.5), Cm(5.5), Cm(5.0), Cm(1.0), Cm(2.0)]

    if idx == 0:
        for ci, h in enumerate(cols_c):
            set_cell(t.cell(0,ci), h, bold=True, size=Pt(9)); t.cell(0,ci).width = wc[ci]; shade(t.cell(0,ci))
        for ci, val in enumerate(row_data[:6]):
            set_cell(t.cell(1,ci), val, size=Pt(8.5)); t.cell(1,ci).width = wc[ci]
    else:
        for ci, val in enumerate(row_data[:6]):
            set_cell(t.cell(0,ci), val, bold=(ci==0), size=Pt(8.5)); t.cell(0,ci).width = wc[ci]

    # Detail row: merge all cells
    detail_cell = t.cell(1,0)
    for ci in range(1, 6):
        detail_cell.merge(t.cell(1,ci))

    p = detail_cell.paragraphs[0]; p.clear()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('Issue: ' + row_data[6])
    run.font.name = 'Microsoft YaHei'; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    rPr = run._element.get_or_add_rPr()
    rf = etree.SubElement(rPr, qn('w:rFonts')); rf.set(qn('w:eastAsia'), 'Microsoft YaHei')
    shade(detail_cell, 'F9FAFB')
    doc.add_paragraph()

# Use simple ASCII name to avoid encoding issues
out_path = os.path.join(desktop, 'StyleRef_v3.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
