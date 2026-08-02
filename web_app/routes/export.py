"""
导出路由 — MD / TXT / HTML / PDF / Word / Excel
专业排版引擎：统一16cm表宽、比例列宽、8pt/9pt分级字体
"""
import re, os, glob
from io import BytesIO
from urllib.parse import quote

from flask import Blueprint, jsonify, request, send_file, Response
from lxml import etree

from database import db
from models.report import Report

bp = Blueprint('export', __name__, url_prefix='/api/export')

PAGE_W = 16.0  # cm, total table width
FONT = 'Microsoft YaHei'


def _filename_docx(report):
    """Word下载文件名 = 举证PPT原名 + 分析报告.docx"""
    ppt = (report.ppt_filename or 'report').strip()
    # Strip .pptx extension if present
    if ppt.lower().endswith('.pptx'):
        ppt = ppt[:-5]
    # Keep only safe characters
    safe = re.sub(r'[\\/:*?"<>|]', '_', ppt)
    return f'{safe}分析报告.docx'


def _filename(report, ext):
    safe = (report.ppt_filename or 'report').strip()
    if safe.lower().endswith('.pptx'):
        safe = safe[:-5]
    safe = re.sub(r'[\\/:*?"<>|]', '_', safe)
    return f'{safe}_report.{ext}'


# ================================================================
# MD / TXT
# ================================================================
@bp.route('/report/<int:report_id>/md', methods=['GET'])
def export_md(report_id):
    r = Report.query.get_or_404(report_id)
    if not r.raw_markdown: return jsonify({'error': 'empty'}), 404
    buf = BytesIO(r.raw_markdown.encode('utf-8')); buf.seek(0)
    return send_file(buf, mimetype='text/markdown; charset=utf-8', as_attachment=True, download_name=_filename(r, 'md'))

@bp.route('/report/<int:report_id>/txt', methods=['GET'])
def export_txt(report_id):
    r = Report.query.get_or_404(report_id)
    if not r.raw_markdown: return jsonify({'error': 'empty'}), 404
    txt = _clean_text(r.raw_markdown)
    buf = BytesIO(txt.encode('utf-8')); buf.seek(0)
    return send_file(buf, mimetype='text/plain; charset=utf-8', as_attachment=True, download_name=_filename(r, 'txt'))

@bp.route('/report/<int:report_id>/html', methods=['GET'])
def export_html(report_id):
    r = Report.query.get_or_404(report_id)
    if not r.raw_markdown: return jsonify({'error': 'empty'}), 404
    return Response(_md_to_html(r.raw_markdown, r.employee_name or ''), mimetype='text/html; charset=utf-8')


# ================================================================
# PDF (fpdf2)
# ================================================================
@bp.route('/report/<int:report_id>/pdf', methods=['GET'])
def export_pdf(report_id):
    r = Report.query.get_or_404(report_id)
    if not r.raw_markdown: return jsonify({'error': 'empty'}), 404
    try:
        from fpdf import FPDF
        pdf = FPDF(); pdf.add_page()
        regular_font, _ = _find_font()
        fn = 'CJK' if regular_font else 'Helvetica'
        if regular_font: pdf.add_font(fn, '', regular_font)

        sections = _parse_sections(r.raw_markdown)
        for sec in sections:
            t = sec['type']
            if t in ('h1', 'h2', 'h3', 'h4'):
                sizes = {'h1': 18, 'h2': 14, 'h3': 12, 'h4': 11}
                pdf.set_font(fn, '', sizes.get(t, 12))
                pdf.ln(4)
                pdf.multi_cell(0, sizes.get(t, 12) * 0.55, _clean_text(sec.get('text', '')))
                if t == 'h1':
                    pdf.set_draw_color(79, 70, 229); pdf.set_line_width(0.8)
                    pdf.line(pdf.l_margin, pdf.get_y() + 2, pdf.w - pdf.r_margin, pdf.get_y() + 2)
                pdf.ln(4)
            elif t == 'blockquote':
                pdf.set_font(fn, '', 9); pdf.set_fill_color(238, 242, 255)
                indent = 6; bw = pdf.w - pdf.l_margin - pdf.r_margin - indent
                for line in sec.get('lines', []):
                    pdf.set_x(pdf.l_margin + indent)
                    pdf.cell(bw, 5, _clean_text(line), new_x='LMARGIN', new_y='NEXT', fill=True)
                pdf.set_x(pdf.l_margin); pdf.ln(2)
            elif t == 'table':
                _pdf_table(pdf, fn, sec)
            elif t == 'hr':
                pdf.set_draw_color(200, 200, 200)
                pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y()); pdf.ln(3)
            elif t == 'list':
                pdf.set_font(fn, '', 10)
                lw = pdf.w - pdf.l_margin - pdf.r_margin - 10
                for line in sec.get('lines', []):
                    clean = _clean_text(line)
                    if clean.startswith('[  ]') or clean.startswith('[]'): clean = '☐ ' + clean[4:] if len(clean) > 4 else clean
                    elif clean.startswith('[x]') or clean.startswith('[X]') or clean.startswith('[OK]'): clean = '☑ ' + clean[4:] if len(clean) > 4 else clean
                    pdf.set_x(pdf.l_margin + 10); pdf.cell(lw, 5, '- ' + clean, new_x='LMARGIN', new_y='NEXT')
                pdf.ln(1)
            elif t == 'para':
                pdf.set_font(fn, '', 10)
                text = _clean_text('\n'.join(sec.get('lines', [])))
                if text.strip(): pdf.multi_cell(0, 5.5, text); pdf.ln(1)

        buf = BytesIO(); pdf.output(buf); buf.seek(0)
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name=_filename(r, 'pdf'))
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': f'PDF: {e}'}), 500


def _pdf_table(pdf, fn, sec):
    """Render a clean table in PDF with proportional column widths."""
    rows = sec.get('rows', [])
    if not rows: return
    cells_list = [[_clean_text(c.strip()) for c in r.split('|') if c.strip()] for r in rows if '|' in r and not r.strip().startswith('|---')]
    if not cells_list: return
    # Proportional widths based on content length
    ncols = max(len(c) for c in cells_list)
    avail = pdf.w - pdf.l_margin - pdf.r_margin
    col_maxes = []
    for ci in range(ncols):
        mx = max((len(cells_list[r][ci]) if ci < len(cells_list[r]) else 0) for r in range(len(cells_list)))
        col_maxes.append(mx if mx > 0 else 1)
    total_mx = sum(col_maxes)
    widths = [max(avail * m / total_mx, avail * 0.07) for m in col_maxes]
    # Rescale
    scale = avail / sum(widths); widths = [w * scale for w in widths]
    # Render header
    pdf.set_fill_color(238, 242, 255); pdf.set_font(pdf.font_family, '', 7)
    for ci, ct in enumerate(cells_list[0]):
        if ci < ncols: pdf.cell(widths[ci], 6, ct[:80], border=1, fill=True)
    pdf.ln()
    # Data rows
    for row in cells_list[1:]:
        for ci, ct in enumerate(row):
            if ci < ncols: pdf.set_font(pdf.font_family, '', 7); pdf.cell(widths[ci], 5.5, ct[:80], border=1)
        pdf.ln()
        if pdf.get_y() > pdf.h - 30: pdf.add_page()
    pdf.set_x(pdf.l_margin); pdf.ln(3)


# ================================================================
# Word (python-docx) — Professional layout
# ================================================================
def markdown_to_docx(md_text: str):
    """将报告Markdown渲染为Word文档对象（可 doc.save(path) 保存）"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn

    doc = Document()
    st = doc.styles['Normal']; st.font.name = FONT; st.font.size = Pt(10)
    rPr = st.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), FONT)
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    for sec in _parse_sections(md_text):
        _write_docx_section(doc, sec)
    return doc


@bp.route('/report/<int:report_id>/docx', methods=['GET'])
def export_docx(report_id):
    r = Report.query.get_or_404(report_id)
    if not r.raw_markdown: return jsonify({'error': 'empty'}), 404
    try:
        doc = markdown_to_docx(r.raw_markdown)
        buf = BytesIO(); doc.save(buf); buf.seek(0)
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=_filename_docx(r))
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': f'Word: {e}'}), 500


def _write_docx_section(doc, sec):
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    t = sec['type']

    if t == 'h1':
        h = doc.add_heading(_clean_text(sec['text']), level=1)
        _font(h.runs[0], Pt(22), RGBColor(0x4F, 0x46, 0xE5))
    elif t == 'h2':
        h = doc.add_heading(_clean_text(sec['text']), level=2)
        _font(h.runs[0], Pt(15), RGBColor(0x1F, 0x29, 0x37))
    elif t == 'h3':
        h = doc.add_heading(_clean_text(sec['text']), level=3)
        _font(h.runs[0], Pt(13))
    elif t == 'h4':
        h = doc.add_heading(_clean_text(sec['text']), level=4)
        _font(h.runs[0], Pt(11))
    elif t == 'blockquote':
        for line in sec['lines']:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1)
            r = p.add_run(_clean_text(line))
            r.font.color.rgb = RGBColor(0x37, 0x41, 0x51); _font(r, Pt(9.5))
    elif t == 'table':
        _docx_table(doc, sec)
    elif t == 'hr':
        p = doc.add_paragraph(); r = p.add_run('─' * 60)
        r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC); _font(r, Pt(8))
    elif t == 'list':
        for line in sec['lines']:
            clean = _clean_text(line)
            if clean.startswith('[  ]') or clean.startswith('[]'): clean = '☐ ' + clean[4:] if len(clean) > 4 else clean
            elif clean.startswith('[x]') or clean.startswith('[X]') or clean.startswith('[OK]'): clean = '☑ ' + clean[4:] if len(clean) > 4 else clean
            p = doc.add_paragraph(clean, style='List Bullet')
            for r in p.runs: _font(r, Pt(10))
    elif t == 'para':
        text = _clean_text('\n'.join(sec['lines']))
        if text.strip():
            p = doc.add_paragraph(text)
            if p.runs: _font(p.runs[0], Pt(10.5))


def _docx_table(doc, sec):
    """Build a professional Word table — unified 16cm width, content-weighted columns.
    文字多的列分配更大列宽，文字少的列（序号/得分/标记等）压缩到最小宽度。"""
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    rows = sec.get('rows', [])
    if len(rows) < 2:
        for r in rows: doc.add_paragraph(_clean_text(r))
        return
    cells_list = [[_clean_text(c.strip()) for c in r.split('|') if c.strip()] for r in rows if '|' in r and not r.strip().startswith('|---')]
    if not cells_list: return
    ncols = max(len(c) for c in cells_list)

    # 列宽评分：以列内容总量为主（60%平均长度 + 40%最长单元格），
    # 使文字多的列显著更宽，仅有短文本/标记的列自然变窄
    MIN_W = 1.4  # cm，短列（序号/得分/✅❌等）最小宽度
    col_scores = []
    for ci in range(ncols):
        lens = [(len(cells_list[r][ci]) if ci < len(cells_list[r]) else 0) for r in range(len(cells_list))]
        avg = sum(lens) / len(lens)
        mx = max(lens)
        col_scores.append(max(0.6 * avg + 0.4 * mx, 1.0))

    # 按评分比例分宽，低于最小宽度的列固定为最小宽度并把余量重新分给长文本列
    fixed = {}
    new_widths = {}
    while True:
        remaining = PAGE_W - MIN_W * len(fixed)
        free = {i: col_scores[i] for i in range(ncols) if i not in fixed}
        if not free:
            break
        total = sum(free.values()) or 1
        new_widths = {i: remaining * s / total for i, s in free.items()}
        under = [i for i, w in new_widths.items() if w < MIN_W]
        if not under:
            break
        for i in under:
            fixed[i] = MIN_W
    widths = [MIN_W if i in fixed else new_widths.get(i, MIN_W) for i in range(ncols)]

    # Choose font size: complex tables (>=5 cols or total chars > 300) use 8pt, simple tables use 9pt
    total_chars = sum(sum(len(c) for c in row) for row in cells_list)
    fs = Pt(8) if (ncols >= 5 or total_chars > 300) else Pt(9)

    table = doc.add_table(rows=len(cells_list), cols=ncols)
    table.style = 'Table Grid'; table.alignment = 1; table.autofit = False
    # 写入网格列宽定义，保证Word按指定列宽渲染
    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for gc, w in zip(tblGrid.findall(qn('w:gridCol')), widths):
            gc.set(qn('w:w'), str(int(w * 567)))  # cm → twips (1cm ≈ 567 twips)
    for ri, row in enumerate(cells_list):
        for ci, text in enumerate(row):
            if ci >= ncols: break
            cell = table.cell(ri, ci); cell.width = Cm(widths[ci])
            p = cell.paragraphs[0]; p.clear()
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(text); _font(run, fs)
            if ri == 0:  # Header row
                run.bold = True
                shd = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
                shd.set(qn('w:fill'), 'EEF2FF'); shd.set(qn('w:val'), 'clear')


# ================================================================
# Excel
# ================================================================
@bp.route('/reports/xlsx', methods=['GET'])
def export_xlsx():
    ids_str = request.args.get('ids', '')
    ids = [int(x.strip()) for x in ids_str.split(',') if x.strip()] if ids_str else []
    q = Report.query.filter(Report.id.in_(ids)) if ids else Report.query.filter(Report.status == 'final').order_by(Report.created_at.desc()).limit(500)
    reports = q.all()
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = '汇总'
    hdrs = ['序号', '姓名', '部门', '岗位', '级别', '总分/16', '结论', '日期']
    hf = PatternFill(start_color='4F46E5', end_color='4F46E5', fill_type='solid')
    for ci, h in enumerate(hdrs, 1):
        c = ws.cell(row=1, column=ci, value=h)
        c.font = Font(bold=True, color='FFFFFF'); c.fill = hf; c.alignment = Alignment(horizontal='center')
    cmap = {'pass': '通过', 'conditional': '有条件通过', 'fail': '不通过'}
    for ri, r in enumerate(reports, 1):
        for ci, v in enumerate([ri, r.employee_name, r.employee_department, r.applied_position, r.applied_level, r.total_score, cmap.get(r.conclusion, str(r.conclusion)), r.created_at.strftime('%Y-%m-%d') if r.created_at else ''], 1):
            ws.cell(row=ri + 1, column=ci, value=v).alignment = Alignment(horizontal='center')
    for ci in range(1, len(hdrs) + 1): ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = 15
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', as_attachment=True, download_name='reports.xlsx')


# ============================================================================
# Shared: Font helper
# ============================================================================
def _font(run, size, color=None):
    from docx.oxml.ns import qn
    run.font.name = FONT; run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), FONT)
    if color: run.font.color.rgb = color


# ============================================================================
# Shared: Find CJK font for PDF
# ============================================================================
def _find_font():
    candidates = ['C:/Windows/Fonts/msyh.ttc', 'C:/Windows/Fonts/simsun.ttc', 'C:/Windows/Fonts/simhei.ttf']
    for p in candidates:
        if os.path.exists(p): return p, None
    for pat in ['C:/Windows/Fonts/msyh*', 'C:/Windows/Fonts/sim*']:
        m = glob.glob(pat)
        if m: return m[0], None
    return None, None


# ============================================================================
# Shared: Markdown → section list
# ============================================================================
def _parse_sections(md_text: str) -> list:
    lines = md_text.split('\n')
    sections = []
    cur = {'type': 'para', 'lines': []}
    in_table = False; table_rows = []

    def flush():
        nonlocal cur, in_table, table_rows
        if in_table and table_rows:
            dr = [r for r in table_rows if '|' in r and not r.strip().startswith('|---')]
            if len(dr) >= 2: sections.append({'type': 'table', 'rows': dr})
            elif table_rows:
                for r in table_rows: sections.append({'type': 'para', 'lines': [r]})
            table_rows = []; in_table = False
        if cur.get('lines') and len(cur['lines']) > 0: sections.append(cur)
        elif cur.get('text'): sections.append(cur)
        cur = {'type': 'para', 'lines': []}

    for line in lines:
        s = line.strip()
        m = re.match(r'^(#{1,4})\s+(.+)', s)
        if m:
            flush(); cur = {'type': f'h{len(m.group(1))}', 'text': m.group(2)}; flush(); continue
        if s == '---': flush(); sections.append({'type': 'hr'}); continue
        if s.startswith('> '):
            if cur['type'] != 'blockquote': flush(); cur = {'type': 'blockquote', 'lines': []}
            cur['lines'].append(re.sub(r'^>\s?', '', s)); continue
        if '|' in s and s.count('|') >= 2:
            if not in_table: flush(); in_table = True
            table_rows.append(s); continue
        if in_table: flush(); in_table = False
        if not s: flush(); continue
        if re.match(r'^-\s', s):
            if cur['type'] != 'list': flush(); cur = {'type': 'list', 'lines': []}
            cur['lines'].append(re.sub(r'^-\s?', '', s)); continue
        if cur['type'] == 'para': cur['lines'].append(s)
        else: flush(); cur = {'type': 'para', 'lines': [s]}
    flush()
    return sections


# ============================================================================
# Shared: Text cleaner — strip **, <br>, HTML tags
# ============================================================================
def _clean_text(text: str) -> str:
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.I)
    text = re.sub(r'</?p>', '', text, flags=re.I)
    text = re.sub(r'</?div[^>]*>', '', text, flags=re.I)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = text.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ============================================================================
# HTML preview
# ============================================================================
def _md_to_html(md_text, title):
    h = md_text
    h = re.sub(r'^#### (.+)$', r'<h4>\1</h4>', h, flags=re.M)
    h = re.sub(r'^### (.+)$', r'<h3>\1</h3>', h, flags=re.M)
    h = re.sub(r'^## (.+)$', r'<h2>\1</h2>', h, flags=re.M)
    h = re.sub(r'^# (.+)$', r'<h1>\1</h1>', h, flags=re.M)
    h = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', h)
    h = re.sub(r'^> (.+)$', r'<blockquote>\1</blockquote>', h, flags=re.M)
    h = h.replace('</blockquote>\n<blockquote>', '<br>')
    h = re.sub(r'^---$', '<hr>', h, flags=re.M)
    h = re.sub(r'^- (.+)$', r'<li>\1</li>', h, flags=re.M)
    h = re.sub(r'(<li>.*</li>\n?)+', r'<ul>\g<0></ul>', h)
    h = h.replace('- [ ] ', '☐ ').replace('- [x] ', '☑ ')
    h = re.sub(r'(\|.+\|\n)+', _html_table, h)
    lines = h.split('\n'); result = []
    for l in lines:
        s = l.strip()
        if not s: result.append('<br>')
        elif s.startswith('<'): result.append(s)
        else: result.append(f'<p>{s}</p>')
    body = '\n'.join(result)
    return f'<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8"><title>{title}</title><style>@page{{size:A4;margin:18mm 15mm}}body{{font-family:"Microsoft YaHei","PingFang SC",sans-serif;max-width:900px;margin:0 auto;padding:20px;color:#1f2937;line-height:1.8;font-size:13px}}h1{{font-size:22px;border-bottom:3px solid #4F46E5;padding-bottom:10px}}h2{{font-size:17px;margin:20px 0 10px}}h3{{font-size:14px}}table{{width:100%;border-collapse:collapse;margin:10px 0 18px;font-size:12px;page-break-inside:avoid}}th{{background:#EEF2FF;color:#4F46E5;padding:6px 8px;text-align:left;border-bottom:2px solid #4F46E5;font-weight:700}}td{{padding:5px 8px;border-bottom:1px solid #E5E7EB}}blockquote{{border-left:4px solid #4F46E5;padding:8px 14px;background:#EEF2FF;margin:10px 0}}@media print{{body{{font-size:11px}}table{{font-size:10px}}}}</style></head><body>{body}</body></html>'

def _html_table(m):
    rows = m.group(0).strip().split('\n')
    dr = [r for r in rows if '|' in r and not r.strip().startswith('|---')]
    if len(dr) < 2: return m.group(0)
    html = '<table>'
    for i, row in enumerate(dr):
        cells = [c.strip() for c in row.split('|') if c.strip()]
        tag = 'th' if i == 0 else 'td'
        html += '<tr>' + ''.join(f'<{tag}>{c}</{tag}>' for c in cells) + '</tr>'
    return html + '</table>'
